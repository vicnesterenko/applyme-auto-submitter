import json
from pathlib import Path
from docx import Document

from resources.values import PROFILE_DOCX, RESUME_DOCX, DEFAULT_RESUME_PDF
from docx_to_pdf import build_resume_pdf


def _docx_text(path: str) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def _extract_profile_json(text: str) -> dict:
    marker = "profile creation"
    idx = text.lower().find(marker)
    chunk = text[idx:] if idx >= 0 else text
    rq_idx = chunk.lower().find("rq")
    chunk = chunk[rq_idx + 2:] if rq_idx >= 0 else chunk
    start = chunk.find("{")
    if start < 0:
        raise ValueError("Profile JSON not found in profile_candidate.docx")
    depth = 0
    for i, ch in enumerate(chunk[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return json.loads(chunk[start: i + 1])
    raise ValueError("Invalid profile JSON in profile_candidate.docx")


def parse_profile_docx(path: str | None = None) -> dict:
    text = _docx_text(path or PROFILE_DOCX)
    return _extract_profile_json(text)


def parse_resume_docx(path: str | None = None) -> str:
    return _docx_text(path or RESUME_DOCX)


def profile_to_candidate(profile: dict, resume_text: str = "") -> dict:
    personal = profile.get("personal_information", {})
    misc = profile.get("miscellaneous", {})
    work = profile.get("work_experience") or []
    skills = profile.get("skills") or []

    industries = profile.get("industries_experience") or []
    wanted_work_type = profile.get("wanted_work_type") or []
    wanted_job_type = profile.get("wanted_job_type") or []
    job_titles = [j.get("job_title", "") for j in profile.get("job_titles", []) if j.get("job_title")]

    first = personal.get("first_name", "")
    last = personal.get("last_name", "")
    phone = personal.get("phone_number", "")
    cc = personal.get("country_code", "")
    if cc and phone and not phone.startswith("+"):
        phone = f"{cc}{phone}"

    current_job = next((w for w in work if w.get("currently_work_here")), work[0] if work else {})
    org = current_job.get("company", "")

    experience_bits = []
    for job in work:
        title = job.get("job_title", "")
        company = job.get("company", "")
        desc = job.get("description", "")
        if title or company:
            experience_bits.append(f"{title} at {company}. {desc}".strip())
    experience = " ".join(experience_bits) or resume_text.strip()

    skill_names = [s.get("skill", "") for s in skills if s.get("skill")]
    skills_str = ", ".join(skill_names)

    city = personal.get("city", "")
    state = personal.get("state", "")
    country = personal.get("country", "")
    location = personal.get("address") or ", ".join(x for x in [city, state, country] if x)

    return {
        "name": f"{first} {last}".strip(),
        "email": personal.get("email", ""),
        "phone": phone,
        "org": org,
        "location": location,
        "linkedin": misc.get("linkedin_url", ""),
        "github": "",
        "portfolio": "",
        "experience": experience,
        "skills": skills_str,
        "summary": resume_text.strip(),
        "gender": personal.get("gender", ""),
        "total_experience": misc.get("total_experience", ""),
        "salary": misc.get("expected_salary_amount", ""),
        "requires_sponsorship": profile.get("requires_sponsorship", False),
        "authorised_to_work": misc.get("authorised_to_work", []),
        "willing_to_relocate": misc.get("willing_to_relocate", "No"),
        "industries_experience": industries,
        "target_job_titles": job_titles,
        "wanted_work_type": wanted_work_type,
        "wanted_job_type": wanted_job_type,
        "auto_apply": profile.get("auto_apply", True)
    }


def ensure_resume_pdf(
        resume_docx: str | None = None,
        output_pdf: str | None = None,
) -> str:
    output_pdf = Path(output_pdf or DEFAULT_RESUME_PDF)
    resume_docx = Path(resume_docx or RESUME_DOCX)

    if output_pdf.exists() and output_pdf.stat().st_size > 1000:
        return str(output_pdf.resolve())

    text = parse_resume_docx(str(resume_docx))
    return build_resume_pdf(text, output_pdf)


def load_profile_file(path: str | Path) -> dict:
    p = Path(path)
    if p.suffix.lower() == ".json":
        return json.loads(p.read_text(encoding="utf-8"))
    return parse_profile_docx(str(p))


def resume_text_from_file(path: str | Path) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _docx_text(str(p))
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")
    return ""


def resume_pdf_from_file(path: str | Path, output_pdf: Path) -> str:
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return str(p.resolve())
    return build_resume_pdf(resume_text_from_file(p), output_pdf)


def load_candidate_data() -> tuple[dict, str]:
    profile = parse_profile_docx()
    resume_text = parse_resume_docx()
    candidate_data = profile_to_candidate(profile, resume_text)
    resume_path_data = ensure_resume_pdf()
    return candidate_data, resume_path_data


if __name__ == "__main__":
    if DEFAULT_RESUME_PDF.exists():
        DEFAULT_RESUME_PDF.unlink()

    candidate, resume_path = load_candidate_data()
    print(json.dumps(candidate, ensure_ascii=False, indent=2))
    print(f"\nResume PDF successfully created at: {resume_path}")

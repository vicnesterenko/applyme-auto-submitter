"""Parsing of candidate profile / resume and resume-PDF generation.

This module contains no network or automation logic — only local file parsing
and PDF rendering. It is the consolidation of the previous duplicated
``applyme_bot.py`` / ``resourses/profile_parser.py`` implementations.
"""

from __future__ import annotations

import json
import textwrap
from pathlib import Path

from docx import Document
from fpdf import FPDF

from .sample_data import RESOURCES_DIR

PROFILE_DOCX = RESOURCES_DIR / "profile_candidate.docx"
RESUME_DOCX = RESOURCES_DIR / "resume.docx"
DEFAULT_RESUME_PDF = RESOURCES_DIR.parent / "resume.pdf"

# Minimum byte size for an existing resume PDF to be considered usable.
_MIN_PDF_BYTES = 5000
_WRAP_WIDTH = 90


def _docx_text(path: str | Path) -> str:
    """Extract visible text (paragraphs + table cells) from a .docx file."""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return "\n".join(parts)


def _extract_profile_json(text: str) -> dict:
    """Pull the first balanced ``{...}`` JSON object out of the profile text."""
    marker = "profile creation"
    idx = text.lower().find(marker)
    chunk = text[idx:] if idx >= 0 else text
    rq_idx = chunk.lower().find("rq")
    chunk = chunk[rq_idx + 2:] if rq_idx >= 0 else chunk

    start = chunk.find("{")
    if start < 0:
        raise ValueError("JSON профілю не знайдено в profile_candidate.docx")

    depth = 0
    for i, ch in enumerate(chunk[start:], start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return json.loads(chunk[start:i + 1])
    raise ValueError("Невалідний JSON профілю в profile_candidate.docx")


def parse_profile_docx(path: str | Path | None = None) -> dict:
    """Parse a candidate profile embedded as JSON inside a .docx file."""
    return _extract_profile_json(_docx_text(path or PROFILE_DOCX))


def parse_resume_docx(path: str | Path | None = None) -> str:
    """Return the plain text of a resume .docx file."""
    return _docx_text(path or RESUME_DOCX)


def profile_to_candidate(profile: dict, resume_text: str = "") -> dict:
    """Map a raw profile dict to the normalised candidate dict used downstream.

    Returns a plain ``dict`` (validation happens at the API boundary via
    :class:`applyme.models.Candidate`).
    """
    personal = profile.get("personal_information", {})
    misc = profile.get("miscellaneous", {})
    work = profile.get("work_experience") or []
    skills = profile.get("skills") or []

    industries = profile.get("industries_experience") or []
    wanted_work_type = profile.get("wanted_work_type") or []
    wanted_job_type = profile.get("wanted_job_type") or []
    job_titles = [
        j.get("job_title", "")
        for j in profile.get("job_titles", [])
        if j.get("job_title")
    ]

    first = personal.get("first_name", "")
    last = personal.get("last_name", "")
    phone = personal.get("phone_number", "")
    cc = personal.get("country_code", "")
    if cc and phone and not phone.startswith("+"):
        phone = f"{cc}{phone}"

    current_job = next(
        (w for w in work if w.get("currently_work_here")),
        work[0] if work else {},
    )
    org = current_job.get("company", "")

    experience_bits = []
    for job in work:
        title = job.get("job_title", "")
        company = job.get("company", "")
        desc = job.get("description", "")
        if title or company:
            experience_bits.append(f"{title} at {company}. {desc}".strip())
    experience = " ".join(experience_bits) or resume_text.strip()

    skill_names = [s.get("skill", "") for s in skills if s.get("skill")]
    skills_str = ", ".join(skill_names)

    city = personal.get("city", "")
    state = personal.get("state", "")
    country = personal.get("country", "")
    location = personal.get("address") or ", ".join(
        x for x in [city, state, country] if x
    )

    return {
        "name": f"{first} {last}".strip(),
        "email": personal.get("email", ""),
        "phone": phone,
        "org": org,
        "location": location,
        "linkedin": misc.get("linkedin_url", ""),
        "github": "",
        "portfolio": "",
        "experience": experience,
        "skills": skills_str,
        "summary": resume_text.strip(),
        "gender": personal.get("gender", ""),
        "total_experience": misc.get("total_experience", ""),
        "salary": misc.get("expected_salary_amount", ""),
        "requires_sponsorship": profile.get("requires_sponsorship", False),
        "authorised_to_work": misc.get("authorised_to_work", []),
        "willing_to_relocate": misc.get("willing_to_relocate", "No"),
        "industries_experience": industries,
        "target_job_titles": job_titles,
        "wanted_work_type": wanted_work_type,
        "wanted_job_type": wanted_job_type,
    }


def build_resume_pdf(text: str, output_pdf: str | Path) -> str:
    """Render plain ``text`` into a simple PDF at ``output_pdf``."""
    output_pdf = Path(output_pdf)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)
    width = pdf.w - pdf.l_margin - pdf.r_margin

    for line in text.splitlines():
        line = line.strip()
        if not line:
            pdf.ln(3)
            continue
        safe = line.encode("latin-1", errors="replace").decode("latin-1")
        for wline in textwrap.wrap(safe, width=_WRAP_WIDTH) or [" "]:
            pdf.multi_cell(width, 5, wline)

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_pdf))
    return str(output_pdf.resolve())


def ensure_resume_pdf(
    resume_docx: str | Path | None = None,
    output_pdf: str | Path | None = None,
) -> str:
    """Generate a resume PDF from a .docx if a usable one does not exist."""
    output_pdf = Path(output_pdf or DEFAULT_RESUME_PDF)
    resume_docx = Path(resume_docx or RESUME_DOCX)

    if output_pdf.exists() and output_pdf.stat().st_size > _MIN_PDF_BYTES:
        return str(output_pdf.resolve())

    return build_resume_pdf(parse_resume_docx(resume_docx), output_pdf)


def load_profile_file(path: str | Path) -> dict:
    """Parse an uploaded candidate profile (.docx or .json)."""
    p = Path(path)
    if p.suffix.lower() == ".json":
        return json.loads(p.read_text(encoding="utf-8"))
    return parse_profile_docx(p)


def resume_text_from_file(path: str | Path) -> str:
    """Extract resume text for the summary field (.docx/.txt/.md; '' for .pdf)."""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".docx":
        return _docx_text(p)
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")
    return ""


def resume_pdf_from_file(path: str | Path, output_pdf: str | Path) -> str:
    """Return a path to a PDF resume: pass through .pdf, else render text."""
    p = Path(path)
    if p.suffix.lower() == ".pdf":
        return str(p.resolve())
    return build_resume_pdf(resume_text_from_file(p), output_pdf)


def load_candidate_data() -> tuple[dict, str]:
    """Convenience loader for the bundled sample profile + resume."""
    profile = parse_profile_docx()
    resume_text = parse_resume_docx()
    candidate = profile_to_candidate(profile, resume_text)
    resume_path = ensure_resume_pdf()
    return candidate, resume_path
